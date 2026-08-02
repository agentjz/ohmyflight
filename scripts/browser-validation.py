from __future__ import annotations

import argparse
import hashlib
import json
import os
import platform
import socket
import statistics
import subprocess
import sys
import tempfile
import threading
import time
from contextlib import contextmanager
from datetime import datetime
from functools import partial
from http.server import SimpleHTTPRequestHandler, ThreadingHTTPServer
from pathlib import Path
from typing import Any, Callable, Iterator

from docx import Document
from openpyxl import Workbook
from playwright.sync_api import Browser, Page, sync_playwright


PROJECT_ROOT = Path(__file__).resolve().parents[1]
DIST_ROOT = Path(os.environ.get("OHMYFLIGHT_DIST_ROOT", PROJECT_ROOT / "dist")).resolve()
DEFAULT_BASELINE = PROJECT_ROOT / "spec" / "dev" / "esm-delivery" / "legacy-performance.json"
HEAVY_PAGES = {
    "training-workbench": "/tool/app/training-workbench/index.html",
    "seasonal-learning": "/tool/app/seasonal-learning/index.html",
    "audit-king": "/tool/app/audit-king/index.html",
    "proof-king": "/tool/app/proof-king/index.html",
}
Operation = Callable[[Page], str]


class QuietHandler(SimpleHTTPRequestHandler):
    def log_message(self, format: str, *args: object) -> None:
        return


def progress(message: str) -> None:
    sys.stderr.write(f"[browser-validation] {message}\n")
    sys.stderr.flush()


@contextmanager
def serve_dist() -> Iterator[str]:
    with socket.socket() as probe:
        probe.bind(("127.0.0.1", 0))
        port = probe.getsockname()[1]
    server = ThreadingHTTPServer(
        ("127.0.0.1", port), partial(QuietHandler, directory=str(DIST_ROOT))
    )
    thread = threading.Thread(target=server.serve_forever, daemon=True)
    thread.start()
    try:
        yield f"http://127.0.0.1:{port}"
    finally:
        server.shutdown()
        server.server_close()
        thread.join(timeout=5)


def create_manual_fixtures(root: Path) -> dict[str, Path]:
    checklist = Document()
    checklist.add_heading("检查单", level=1)
    checklist.add_paragraph("1.1 检查飞行运行程序是否具有完整依据。")
    checklist.add_paragraph("1.2 检查培训记录是否按要求保存。")
    checklist_path = root / "checklist.docx"
    checklist.save(checklist_path)

    audit_manual = Document()
    audit_manual.add_heading("第一章 运行程序", level=1)
    audit_manual.add_paragraph("飞行运行程序应当具有完整依据，并由责任部门定期复核。")
    audit_manual.add_paragraph("培训记录应当按要求保存并能够追溯。")
    audit_manual_path = root / "audit-manual.docx"
    audit_manual.save(audit_manual_path)

    my_manual = Document()
    my_manual.add_heading("第一章 总则", level=1)
    my_manual.add_paragraph("本手册适用于运行管理。")
    my_manual.add_paragraph("记录应保存十二个月。")
    my_manual.add_paragraph("相同锚点内容用于确认顺序。")
    my_manual_path = root / "my-manual.docx"
    my_manual.save(my_manual_path)

    reference_manual = Document()
    reference_manual.add_heading("第一章 总则", level=1)
    reference_manual.add_paragraph("本手册适用于运行管理。")
    reference_manual.add_paragraph("记录应保存二十四个月。")
    reference_manual.add_paragraph("新增的复核要求应当单独记录。")
    reference_manual.add_paragraph("相同锚点内容用于确认顺序。")
    reference_manual_path = root / "reference-manual.docx"
    reference_manual.save(reference_manual_path)
    return {
        "checklist": checklist_path,
        "audit_manual": audit_manual_path,
        "my_manual": my_manual_path,
        "reference_manual": reference_manual_path,
    }


def create_excel_fixtures(root: Path) -> tuple[Path, Path]:
    training = Workbook()
    people = training.active
    people.title = "人员信息表"
    people.append(["员工号", "姓名", "危险品", "航空安保"])
    people.append(["1001", "张三", datetime(2026, 5, 20), datetime(2026, 6, 30)])
    people.append(["1002", "李四", datetime(2026, 5, 31), datetime(2026, 7, 31)])
    headers = ["员工号", "姓名", "项目名称", "培训信息是否录入", "培训开始日期", "培训结束日期", "有效期", "备注"]
    danger = training.create_sheet("危险品")
    danger.append(headers)
    danger.append(["1001", "张三", "危险品", "是", datetime(2026, 4, 15), datetime(2026, 4, 15), "", ""])
    danger.append(["1002", "李四", "危险品", "否", datetime(2026, 5, 10), datetime(2026, 5, 10), "", ""])
    training.create_sheet("航空安保").append(headers)
    training_path = root / "training.xlsx"
    training.save(training_path)

    seasonal = Workbook()
    seasonal.active.title = "换季实际"
    seasonal.active.append(["序号", "员工号", "姓名", "分部", "技术信息", "是否带队", "培训类型", "日期", "期数", "身份"])
    roster = seasonal.create_sheet("换季总名单")
    roster.append(["序号", "员工号", "姓名", "分部", "技术信息", "是否带队", "是否美线带队", "培训类型", "日期", "期数", "身份"])
    sequence = 1
    for technical_info, is_leader, count in (
        ("777:飞行教员A", 1, 6),
        ("777:C类机长", 0, 12),
        ("777:C类副驾驶", 0, 12),
    ):
        for _ in range(count):
            roster.append([
                sequence,
                str(200000 + sequence),
                f"测试人员{sequence}",
                f"{(sequence % 4) + 1}分部",
                technical_info,
                is_leader,
                1 if is_leader and sequence <= 2 else 0,
                "换季学习",
                "",
                "",
                "",
            ])
            sequence += 1
    seasonal_path = root / "seasonal.xlsx"
    seasonal.save(seasonal_path)
    return training_path, seasonal_path


def open_page(browser: Browser, url: str) -> tuple[Any, Page, list[str]]:
    context = browser.new_context()
    page = context.new_page()
    errors: list[str] = []
    page.on("pageerror", lambda error: errors.append(f"pageerror: {error}"))
    page.on(
        "console",
        lambda message: errors.append(f"console.{message.type}: {message.text}")
        if message.type == "error"
        else None,
    )
    page.on(
        "requestfailed",
        lambda request: errors.append(f"requestfailed: {request.url} ({request.failure or 'unknown'})"),
    )
    page.on(
        "response",
        lambda response: errors.append(f"http {response.status}: {response.url}")
        if response.status >= 400
        else None,
    )
    page.goto(url, wait_until="load", timeout=120_000)
    return context, page, errors


def text_hash(page: Page, selectors: list[str]) -> str:
    normalized = "\n".join(
        " ".join(page.locator(selector).inner_text().split()) for selector in selectors
    )
    return hashlib.sha256(normalized.encode("utf-8")).hexdigest()


def reset_file_input(page: Page, selector: str) -> None:
    page.locator(selector).set_input_files([])


def perform_training(page: Page, workbook: Path) -> str:
    page.locator("#statusLine").evaluate("element => { element.textContent = ''; }")
    reset_file_input(page, "#workbookFile")
    page.locator("#workbookFile").set_input_files(str(workbook))
    page.wait_for_function(
        "document.querySelector('#statusLine')?.textContent?.includes('识别完成')",
        timeout=180_000,
    )
    return text_hash(page, ["#workbookOverview", "#workbookHealthPanel", "#scheduleGapSummary", "#statsGrid"])


def perform_seasonal(page: Page, workbook: Path) -> str:
    page.locator("#workspace").evaluate("element => { element.hidden = true; }")
    reset_file_input(page, "#workbookFile")
    page.locator("#workbookFile").set_input_files(str(workbook))
    page.wait_for_function("!document.querySelector('#workspace')?.hidden", timeout=120_000)
    button = page.locator("#balanceButton")
    if button.inner_text().strip() == "均衡负载" and button.is_enabled():
        button.click()
        page.wait_for_function(
            "document.querySelector('#balanceButton')?.textContent?.includes('均衡检查')",
            timeout=120_000,
        )
    return text_hash(page, ["#summaryStrip", "#balanceResultSummary", "#balanceResults"])


def perform_audit(page: Page, fixtures: dict[str, Path]) -> str:
    page.locator("#checklistText").evaluate("element => { element.textContent = ''; }")
    page.locator("#manualList").evaluate("element => { element.textContent = ''; }")
    reset_file_input(page, "#checklistInput")
    reset_file_input(page, "#manualInput")
    page.locator("#checklistInput").set_input_files(str(fixtures["checklist"]))
    page.locator("#manualInput").set_input_files(str(fixtures["audit_manual"]))
    page.wait_for_function(
        "document.querySelector('#checklistText')?.textContent?.includes('检查飞行运行程序')",
        timeout=120_000,
    )
    page.wait_for_function(
        "document.querySelector('#manualList')?.textContent?.includes('audit-manual.docx')",
        timeout=120_000,
    )
    return text_hash(page, ["#checklistText", "#manualList", "#statusBar"])


def perform_proof(page: Page, fixtures: dict[str, Path]) -> str:
    page.locator("#myStatus").evaluate("element => { element.textContent = '未上传'; }")
    page.locator("#referenceStatus").evaluate("element => { element.textContent = '未上传'; }")
    page.locator("#summaryGrid").evaluate("element => { element.replaceChildren(); }")
    reset_file_input(page, "#myInput")
    reset_file_input(page, "#referenceInput")
    page.locator("#myInput").set_input_files(str(fixtures["my_manual"]))
    page.locator("#referenceInput").set_input_files(str(fixtures["reference_manual"]))
    page.wait_for_function(
        "!document.querySelector('#myStatus')?.textContent?.includes('未上传') && "
        "!document.querySelector('#referenceStatus')?.textContent?.includes('未上传')",
        timeout=120_000,
    )
    page.locator("#compareButton").click()
    page.wait_for_function("document.querySelector('#summaryGrid')?.children.length > 0", timeout=120_000)
    return text_hash(page, ["#summaryGrid", "#resultCount", "#workspaceMessage"])


def summarize(values: list[float]) -> dict[str, float]:
    ordered = sorted(values)
    p95_index = min(len(ordered) - 1, round((len(ordered) - 1) * 0.95))
    return {
        "median": round(statistics.median(values), 2),
        "p95": round(ordered[p95_index], 2),
        "minimum": round(min(values), 2),
        "maximum": round(max(values), 2),
    }


def startup_metrics(page: Page, wall_ms: float) -> dict[str, float]:
    result = page.evaluate(
        """() => {
          const nav = performance.getEntriesByType('navigation')[0];
          const resources = performance.getEntriesByType('resource');
          const scripts = resources.filter((entry) => entry.initiatorType === 'script');
          const application = scripts.filter((entry) => {
            const path = new URL(entry.name).pathname;
            return !path.includes('/libs/') && !path.endsWith('/theme.js');
          });
          return {
            navigationMs: nav ? nav.duration : 0,
            domContentLoadedMs: nav ? nav.domContentLoadedEventEnd : 0,
            scriptRequests: scripts.length,
            applicationScriptRequests: application.length,
            transferBytes: resources.reduce((sum, entry) => sum + (entry.transferSize || 0), 0),
            usedJsHeapBytes: performance.memory ? performance.memory.usedJSHeapSize : 0
          };
        }"""
    )
    result["wallMs"] = wall_ms
    return result


def run_startup(browser: Browser, base_url: str, page_path: str) -> tuple[dict[str, float], list[str]]:
    started = time.perf_counter()
    context, page, errors = open_page(browser, f"{base_url}{page_path}")
    metrics = startup_metrics(page, (time.perf_counter() - started) * 1000)
    context.close()
    return metrics, errors


def run_workflow(
    browser: Browser,
    base_url: str,
    page_path: str,
    operation: Operation,
) -> tuple[float, str, list[str]]:
    context, page, errors = open_page(browser, f"{base_url}{page_path}")
    started = time.perf_counter()
    result_hash = operation(page)
    elapsed = (time.perf_counter() - started) * 1000
    context.close()
    return elapsed, result_hash, errors


def run_repeated(
    repetitions: int,
    callback: Callable[[], tuple[float, str, list[str]]],
) -> dict[str, Any]:
    durations: list[float] = []
    hashes: set[str] = set()
    errors: list[str] = []
    for _ in range(repetitions):
        duration, result_hash, run_errors = callback()
        durations.append(duration)
        hashes.add(result_hash)
        errors.extend(run_errors)
    return {
        "durationMs": summarize(durations),
        "resultSha256": sorted(hashes),
        "errors": sorted(set(errors)),
    }


def run_all_pages(browser: Browser, base_url: str) -> list[dict[str, Any]]:
    results: list[dict[str, Any]] = []
    pages = sorted(path.relative_to(DIST_ROOT).as_posix() for path in DIST_ROOT.rglob("*.html"))
    for relative_path in pages:
        context, page, errors = open_page(browser, f"{base_url}/{relative_path}")
        page.wait_for_timeout(250)
        if not page.locator("body").inner_text().strip() and page.locator("body *").count() == 0:
            errors.append("blank page: body has no text or elements")
        results.append({
            "page": relative_path,
            "finalUrl": page.url.removeprefix(f"{base_url}/"),
            "errors": sorted(set(errors)),
        })
        context.close()
    return results


def run_memory_cycles(
    browser: Browser,
    base_url: str,
    operations: dict[str, Operation],
) -> dict[str, Any]:
    results: dict[str, Any] = {}
    for name, operation in operations.items():
        context, page, errors = open_page(browser, f"{base_url}{HEAVY_PAGES[name]}")
        heaps: list[int] = []
        for cycle in range(5):
            progress(f"{name} 内存循环 {cycle + 1}/5")
            operation(page)
            page.request_gc()
            heaps.append(int(page.evaluate("performance.memory ? performance.memory.usedJSHeapSize : 0")))
        context.close()
        results[name] = {
            "usedJsHeapBytes": heaps,
            "growthBytes": heaps[-1] - heaps[0],
            "growthRatio": round(heaps[-1] / heaps[0], 4) if heaps[0] else 0,
            "errors": sorted(set(errors)),
        }
    return results


def validate(result: dict[str, Any], baseline_path: Path | None) -> list[str]:
    failures: list[str] = []
    pages = result.get("pages", [])
    for page in pages:
        if page["errors"]:
            failures.append(f"{page['page']} 页面启动错误：{page['errors']}")
    if pages and len(pages) != 26:
        failures.append(f"构建 HTML 数量为 {len(pages)}，预期 26")

    for name, metrics in result["coldStart"].items():
        if metrics["errors"]:
            failures.append(f"{name} 冷启动错误：{metrics['errors']}")
        if metrics["applicationScriptRequests"]["maximum"] != 1:
            failures.append(f"{name} 应用脚本请求数不是 1")
    for name, workflow in result["workflows"].items():
        if workflow["errors"] or len(workflow["resultSha256"]) != 1:
            failures.append(f"{name} 工作流错误或结果不稳定：{workflow}")
    for name, memory in result["memoryCycles"].items():
        heap = memory["usedJsHeapBytes"]
        allowed_last = max(heap[0] * 1.35, heap[0] + 10 * 1024 * 1024)
        if memory["errors"] or heap[-1] > allowed_last:
            failures.append(f"{name} 内存复核失败：{memory}")

    if baseline_path:
        baseline = json.loads(baseline_path.read_text(encoding="utf-8"))
        if result["method"]["warmups"] < 1 or result["method"]["repetitions"] < 3:
            failures.append("正式性能验收要求预热至少 1 次并采样至少 3 次")
        for name, metrics in result["coldStart"].items():
            median = metrics["wallMs"]["median"]
            baseline_median = baseline["coldStart"][name]["wallMs"]["median"]
            budget = max(baseline_median * 1.25, baseline_median + 100)
            if median > budget:
                failures.append(f"{name} 冷启动中位耗时 {median}ms 超过预算 {round(budget, 2)}ms")
        for name, workflow in result["workflows"].items():
            expected = baseline["workflows"][name]
            if workflow["resultSha256"] != expected["resultSha256"]:
                failures.append(f"{name} 结果哈希变化")
            median = workflow["durationMs"]["median"]
            baseline_median = expected["durationMs"]["median"]
            budget = max(baseline_median * 1.25, baseline_median + 100)
            if median > budget:
                failures.append(f"{name} 中位耗时 {median}ms 超过预算 {round(budget, 2)}ms")
    return failures


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser()
    parser.add_argument("--isolated-build", action="store_true")
    parser.add_argument("--training-workbook", type=Path)
    parser.add_argument("--seasonal-workbook", type=Path)
    parser.add_argument("--synthetic", action="store_true")
    parser.add_argument("--all-pages", action="store_true")
    parser.add_argument("--repetitions", type=int, default=3)
    parser.add_argument("--warmups", type=int, default=1)
    parser.add_argument("--skip-memory", action="store_true")
    parser.add_argument("--baseline", type=Path)
    parser.add_argument("--label", default="current")
    parser.add_argument("--output", type=Path)
    return parser.parse_args()


def run_isolated_build() -> None:
    npm_command = "npm.cmd" if os.name == "nt" else "npm"
    child_arguments = [argument for argument in sys.argv[1:] if argument != "--isolated-build"]
    with tempfile.TemporaryDirectory(prefix="ohmyflight-build-browser-") as output_directory:
        environment = os.environ.copy()
        environment["OHMYFLIGHT_DIST_ROOT"] = output_directory
        subprocess.run(
            [npm_command, "run", "build"],
            cwd=PROJECT_ROOT,
            env=environment,
            check=True,
        )
        completed = subprocess.run(
            [sys.executable, "-X", "utf8", str(Path(__file__).resolve()), *child_arguments],
            cwd=PROJECT_ROOT,
            env=environment,
        )
        if completed.returncode:
            raise SystemExit(completed.returncode)


def main() -> None:
    args = parse_args()
    if args.isolated_build:
        run_isolated_build()
        return
    if not DIST_ROOT.is_dir():
        raise SystemExit("dist/ 不存在，请先运行 npm.cmd run build。")
    if not args.synthetic and (not args.training_workbook or not args.seasonal_workbook):
        raise SystemExit("请传入培训和换季两份 Excel，或使用 --synthetic。")

    with tempfile.TemporaryDirectory(prefix="ohmyflight-browser-") as temp_directory:
        temp_root = Path(temp_directory)
        fixtures = create_manual_fixtures(temp_root)
        if args.synthetic:
            training_workbook, seasonal_workbook = create_excel_fixtures(temp_root)
            input_kind = "generated-xlsx"
        else:
            training_workbook = args.training_workbook
            seasonal_workbook = args.seasonal_workbook
            input_kind = "owner-local"
        assert training_workbook is not None and seasonal_workbook is not None
        for workbook in (training_workbook, seasonal_workbook):
            if not workbook.is_file():
                raise SystemExit(f"本地回放文件不存在：{workbook}")

        operations: dict[str, Operation] = {
            "training-workbench": lambda page: perform_training(page, training_workbook),
            "seasonal-learning": lambda page: perform_seasonal(page, seasonal_workbook),
            "audit-king": lambda page: perform_audit(page, fixtures),
            "proof-king": lambda page: perform_proof(page, fixtures),
        }
        with serve_dist() as base_url, sync_playwright() as playwright:
            browser = playwright.chromium.launch(
                headless=True,
                args=["--enable-precise-memory-info"],
            )
            pages = run_all_pages(browser, base_url) if args.all_pages else []
            for _ in range(args.warmups):
                progress("预热四个重型页面和工作流")
                for name, operation in operations.items():
                    _, startup_errors = run_startup(browser, base_url, HEAVY_PAGES[name])
                    if startup_errors:
                        raise RuntimeError(f"{name} 冷启动预热失败：{startup_errors}")
                    duration, _, errors = run_workflow(browser, base_url, HEAVY_PAGES[name], operation)
                    if errors:
                        raise RuntimeError(f"{name} 工作流预热失败：{errors}; {duration}ms")

            cold_start: dict[str, Any] = {}
            for name, page_path in HEAVY_PAGES.items():
                progress(f"采样冷启动：{name}")
                samples = [run_startup(browser, base_url, page_path) for _ in range(args.repetitions)]
                keys = samples[0][0].keys()
                cold_start[name] = {
                    key: summarize([metrics[key] for metrics, _ in samples]) for key in keys
                }
                cold_start[name]["errors"] = sorted(
                    {error for _, errors in samples for error in errors}
                )

            workflows: dict[str, Any] = {}
            for name, operation in operations.items():
                progress(f"采样工作流：{name}")
                workflows[name] = run_repeated(
                    args.repetitions,
                    lambda name=name, operation=operation: run_workflow(
                        browser, base_url, HEAVY_PAGES[name], operation
                    ),
                )
            memory_cycles = {} if args.skip_memory else run_memory_cycles(browser, base_url, operations)
            browser_version = browser.version
            browser.close()

        result = {
            "schemaVersion": 1,
            "label": args.label,
            "method": {
                "warmups": args.warmups,
                "repetitions": args.repetitions,
                "browserContext": "fresh context per timed sample",
            },
            "environment": {
                "platform": platform.platform(),
                "python": platform.python_version(),
                "playwright": "1.58.0",
                "chromium": browser_version,
                "cpuCount": os.cpu_count(),
            },
            "inputs": {
                "training-workbench": {"kind": input_kind, "bytes": training_workbook.stat().st_size},
                "seasonal-learning": {"kind": input_kind, "bytes": seasonal_workbook.stat().st_size},
                "audit-king": {"kind": "generated-docx"},
                "proof-king": {"kind": "generated-docx-pair"},
            },
            "pages": pages,
            "coldStart": cold_start,
            "workflows": workflows,
            "memoryCycles": memory_cycles,
        }

    baseline_path = args.baseline or (None if args.synthetic else DEFAULT_BASELINE)
    if baseline_path and not baseline_path.is_absolute():
        baseline_path = PROJECT_ROOT / baseline_path
    serialized = json.dumps(result, ensure_ascii=False, indent=2) + "\n"
    if args.output:
        output = args.output if args.output.is_absolute() else PROJECT_ROOT / args.output
        output.parent.mkdir(parents=True, exist_ok=True)
        output.write_text(serialized, encoding="utf-8")
    sys.stdout.write(serialized)
    failures = validate(result, baseline_path)
    if failures:
        raise SystemExit("浏览器验证失败：\n- " + "\n- ".join(failures))


if __name__ == "__main__":
    main()
